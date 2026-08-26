import os
from pathlib import Path
from typing import Dict, Any, List, Tuple
from app.ingestion.ocr.provider import OCRProvider, get_ocr_provider


class DocumentExtractor:
    """
    Unified multi-format document extraction engine supporting PDF, DOCX, and Images.
    Uses PyMuPDF (fitz) and pdfplumber for fast/accurate PDF extraction,
    python-docx for Word files, and OCRProvider for scanned documents / images.
    """

    def __init__(self, ocr_provider: OCRProvider = None):
        self.ocr_provider = ocr_provider or get_ocr_provider()

    def extract(self, file_path: Path) -> Tuple[str, int, Dict[str, Any]]:
        """
        Extracts text, page count, and metadata from the given file.
        Returns: (raw_text, page_count, metadata)
        """
        file_path = Path(file_path)
        suffix = file_path.suffix.lower()

        if suffix == ".pdf":
            return self._extract_pdf(file_path)
        elif suffix in [".docx", ".doc"]:
            return self._extract_docx(file_path)
        elif suffix in [".png", ".jpg", ".jpeg", ".bmp", ".tiff", ".webp"]:
            return self._extract_image(file_path)
        elif suffix in [".txt", ".csv"]:
            return self._extract_text_file(file_path)
        else:
            # Fallback
            return self._extract_generic(file_path)

    def _extract_pdf(self, file_path: Path) -> Tuple[str, int, Dict[str, Any]]:
        text_content: List[str] = []
        page_count = 0
        tables_found = 0

        # Try PyMuPDF first for speed and text streams
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(file_path)
            page_count = len(doc)
            for page_num in range(page_count):
                page = doc.load_page(page_num)
                page_text = page.get_text()
                if page_text.strip():
                    text_content.append(page_text)
            doc.close()
        except Exception:
            pass

        # If PyMuPDF found little or no text (e.g. scanned PDF), try pdfplumber or OCR
        if not "".join(text_content).strip():
            try:
                import pdfplumber
                with pdfplumber.open(file_path) as pdf:
                    page_count = len(pdf.pages)
                    for page in pdf.pages:
                        extracted = page.extract_text()
                        if extracted:
                            text_content.append(extracted)
                        tables = page.extract_tables()
                        if tables:
                            tables_found += len(tables)
                            for table in tables:
                                for row in table:
                                    filtered_row = [str(c) if c is not None else "" for c in row]
                                    text_content.append(" | ".join(filtered_row))
            except Exception:
                pass

        # If still empty, attempt OCR extraction
        if not "".join(text_content).strip():
            ocr_text = self.ocr_provider.extract_text(file_path)
            text_content.append(ocr_text)

        full_text = "\n".join(text_content).strip()
        metadata = {
            "engine": "PyMuPDF/pdfplumber",
            "page_count": max(page_count, 1),
            "tables_found": tables_found,
            "char_count": len(full_text),
        }
        return full_text, max(page_count, 1), metadata

    def _extract_docx(self, file_path: Path) -> Tuple[str, int, Dict[str, Any]]:
        text_content: List[str] = []
        try:
            import docx
            doc = docx.Document(file_path)
            for para in doc.paragraphs:
                if para.text.strip():
                    text_content.append(para.text)
            for table in doc.tables:
                for row in table.rows:
                    row_cells = [c.text.strip() for c in row.cells if c.text.strip()]
                    if row_cells:
                        text_content.append(" | ".join(row_cells))
        except Exception as e:
            text_content.append(f"[DOCX Extraction Fallback: {file_path.name} | {str(e)}]")

        full_text = "\n".join(text_content).strip()
        metadata = {
            "engine": "python-docx",
            "page_count": 1,
            "char_count": len(full_text),
        }
        return full_text, 1, metadata

    def _extract_image(self, file_path: Path) -> Tuple[str, int, Dict[str, Any]]:
        ocr_text = self.ocr_provider.extract_text(file_path)
        metadata = {
            "engine": "OCR",
            "page_count": 1,
            "char_count": len(ocr_text),
        }
        return ocr_text, 1, metadata

    def _extract_text_file(self, file_path: Path) -> Tuple[str, int, Dict[str, Any]]:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            content = f.read()
        return content.strip(), 1, {"engine": "text_reader", "char_count": len(content)}

    def _extract_generic(self, file_path: Path) -> Tuple[str, int, Dict[str, Any]]:
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                content = f.read()
            return content.strip(), 1, {"engine": "generic_reader", "char_count": len(content)}
        except Exception as e:
            return f"[Generic reader error: {str(e)}]", 1, {"engine": "generic_reader", "error": str(e)}
